"""Core functionality for combining Word documents."""
from pathlib import Path
from typing import List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


def combine_documents(
    input_files: List[Path],
    output_path: Path,
    separator: str = 'page_break'
) -> None:
    """
    Combine multiple Word documents into a single document.
    
    Args:
        input_files: List of paths to input .docx files
        output_path: Path for the output combined document
        separator: How to separate documents ('page_break', 'newline', or 'none')
    
    Raises:
        FileNotFoundError: If any input file doesn't exist
        ValueError: If separator is invalid
    """
    if not input_files:
        raise ValueError("At least one input file is required")
    
    # Create a new document for the output
    combined_doc = Document()
    
    for i, input_file in enumerate(input_files):
        if not input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")
        
        # Open each input document
        doc = Document(input_file)
        
        # Add separator before each document (except the first)
        if i > 0:
            if separator == 'page_break':
                # Add a page break
                combined_doc.add_page_break()
            elif separator == 'newline':
                # Add a paragraph break
                combined_doc.add_paragraph()
        
        # Copy all paragraphs from the input document
        for paragraph in doc.paragraphs:
            new_para = combined_doc.add_paragraph()
            
            # Copy paragraph formatting
            new_para.style = paragraph.style
            
            # Copy runs (text with formatting)
            for run in paragraph.runs:
                new_run = new_para.add_run(run.text)
                # Copy run formatting
                if run.bold is not None:
                    new_run.bold = run.bold
                if run.italic is not None:
                    new_run.italic = run.italic
                if run.underline is not None:
                    new_run.underline = run.underline
                if run.font.size is not None:
                    new_run.font.size = run.font.size
                if run.font.name is not None:
                    new_run.font.name = run.font.name
                if run.font.color.rgb is not None:
                    new_run.font.color.rgb = run.font.color.rgb
            
            # Copy paragraph alignment
            if paragraph.alignment is not None:
                new_para.alignment = paragraph.alignment
        
        # Copy tables from the input document
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = combined_doc.add_table(
                rows=len(table.rows),
                cols=len(table.columns)
            )
            
            # Copy table content
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    new_cell = new_table.rows[row_idx].cells[col_idx]
                    new_cell.text = cell.text
                    
                    # Copy cell formatting if possible
                    if cell.paragraphs:
                        for para in cell.paragraphs:
                            new_para = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                            new_para.text = para.text
                            if para.alignment is not None:
                                new_para.alignment = para.alignment
        
        # Add separator after document if needed (for 'none' case, we still want spacing)
        if separator == 'newline' and i < len(input_files) - 1:
            combined_doc.add_paragraph()
    
    # Save the combined document
    combined_doc.save(output_path)
